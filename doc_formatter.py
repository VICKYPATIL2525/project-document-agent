"""
Word Document Formatter - Creates professionally formatted .docx files.
Handles all formatting: title page, headings, code blocks, diagrams placeholders,
tables, page numbers, table of contents, etc.
"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from config import (
    DEFAULT_FONT,
    HEADING1_SIZE,
    HEADING2_SIZE,
    HEADING3_SIZE,
    BODY_SIZE,
    CODE_FONT,
    CODE_SIZE,
    LINE_SPACING,
    PAGE_MARGIN_INCHES,
    OUTPUT_DIR,
)


class DocumentFormatter:
    """Creates professionally formatted Word documents."""

    def __init__(self, filename: str):
        self.doc = Document()
        self.filename = filename
        self.figure_count = 0
        self.table_count = 0
        self._setup_styles()
        self._setup_page_layout()

    def _setup_page_layout(self):
        """Configure page size, margins, etc."""
        for section in self.doc.sections:
            section.page_width = Inches(8.27)   # A4
            section.page_height = Inches(11.69)  # A4
            section.top_margin = Inches(PAGE_MARGIN_INCHES)
            section.bottom_margin = Inches(PAGE_MARGIN_INCHES)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(PAGE_MARGIN_INCHES)

    def _setup_styles(self):
        """Setup custom document styles."""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = DEFAULT_FONT
        font.size = Pt(BODY_SIZE)
        font.color.rgb = RGBColor(0, 0, 0)

        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = LINE_SPACING
        paragraph_format.space_after = Pt(6)

        # Heading 1
        h1 = self.doc.styles["Heading 1"]
        h1.font.name = DEFAULT_FONT
        h1.font.size = Pt(HEADING1_SIZE)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(12)
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Heading 2
        h2 = self.doc.styles["Heading 2"]
        h2.font.name = DEFAULT_FONT
        h2.font.size = Pt(HEADING2_SIZE)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 0, 0)
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(8)

        # Heading 3
        h3 = self.doc.styles["Heading 3"]
        h3.font.name = DEFAULT_FONT
        h3.font.size = Pt(HEADING3_SIZE)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(0, 0, 0)
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(6)

        # Code style
        try:
            code_style = self.doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            code_style = self.doc.styles["Code"]
        code_style.font.name = CODE_FONT
        code_style.font.size = Pt(CODE_SIZE)
        code_style.font.color.rgb = RGBColor(30, 30, 30)
        code_style.paragraph_format.line_spacing = 1.0
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(4)
        code_style.paragraph_format.left_indent = Inches(0.3)

    def add_title_page(
        self,
        project_title: str,
        student_name: str = "Student Name",
        roll_number: str = "Roll Number",
        college_name: str = "College Name",
        department: str = "Department of Computer Science",
        guide_name: str = "Guide Name",
        year: str = "2025-2026",
    ):
        """Add a formatted title page."""
        # Add some spacing at top
        for _ in range(4):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

        # College name
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(college_name.upper())
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = DEFAULT_FONT

        # Department
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(department)
        run.font.size = Pt(14)
        run.font.name = DEFAULT_FONT

        # Spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Project title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("A Project Report On")
        run.font.size = Pt(14)
        run.font.name = DEFAULT_FONT

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'"{project_title}"')
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = DEFAULT_FONT
        run.font.color.rgb = RGBColor(0, 0, 120)

        # Spacing
        self.doc.add_paragraph()

        # Submitted by
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Submitted By")
        run.font.size = Pt(12)
        run.font.name = DEFAULT_FONT

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(student_name)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = DEFAULT_FONT

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Roll No: {roll_number}")
        run.font.size = Pt(12)
        run.font.name = DEFAULT_FONT

        # Spacing
        self.doc.add_paragraph()

        # Under the guidance of
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Under the Guidance of")
        run.font.size = Pt(12)
        run.font.name = DEFAULT_FONT

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(guide_name)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = DEFAULT_FONT

        # Spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Academic Year
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Academic Year: {year}")
        run.font.size = Pt(14)
        run.font.name = DEFAULT_FONT
        run.font.bold = True

        self.doc.add_page_break()

    def add_certificate_page(
        self,
        project_title: str,
        student_name: str = "Student Name",
        roll_number: str = "Roll Number",
        college_name: str = "College Name",
        department: str = "Department of Computer Science",
        guide_name: str = "Guide Name",
    ):
        """Add a certificate page."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CERTIFICATE")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = DEFAULT_FONT
        run.font.underline = True

        self.doc.add_paragraph()

        text = (
            f"This is to certify that the project entitled "
            f'"{project_title}" is a bonafide work carried out by '
            f"{student_name} (Roll No: {roll_number}), "
            f"a student of {department}, {college_name}, "
            f"in partial fulfillment of the requirements for the award of the degree."
        )
        p = self.doc.add_paragraph(text)
        p.paragraph_format.line_spacing = LINE_SPACING
        for run in p.runs:
            run.font.size = Pt(BODY_SIZE)
            run.font.name = DEFAULT_FONT

        for _ in range(4):
            self.doc.add_paragraph()

        # Signature section
        table = self.doc.add_table(rows=2, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cells = table.rows[0].cells
        for cell in cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = cell.paragraphs[0]
            run = p.add_run("_________________")
            run.font.size = Pt(BODY_SIZE)
            run.font.name = DEFAULT_FONT

        cells = table.rows[1].cells
        labels = ["Project Guide", "Head of Department", "External Examiner"]
        for cell, label in zip(cells, labels):
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(label)
            run.font.size = Pt(BODY_SIZE)
            run.font.bold = True
            run.font.name = DEFAULT_FONT

        self.doc.add_page_break()

    def add_acknowledgment(self, content: str):
        """Add acknowledgment page."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ACKNOWLEDGMENT")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = DEFAULT_FONT
        run.font.underline = True

        self.doc.add_paragraph()
        self._add_formatted_text(content)
        self.doc.add_page_break()

    def add_abstract(self, content: str):
        """Add abstract page."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ABSTRACT")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = DEFAULT_FONT
        run.font.underline = True

        self.doc.add_paragraph()
        self._add_formatted_text(content)
        self.doc.add_page_break()

    def add_table_of_contents(self):
        """Add a Table of Contents field (auto-updates in Word)."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("TABLE OF CONTENTS")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = DEFAULT_FONT
        run.font.underline = True

        self.doc.add_paragraph()

        # Add TOC field code
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar)

        run = paragraph.add_run()
        instrText = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'
        )
        run._r.append(instrText)

        run = paragraph.add_run()
        fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run._r.append(fldChar)

        run = paragraph.add_run("Right-click and select 'Update Field' to generate Table of Contents")
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.size = Pt(11)

        run = paragraph.add_run()
        fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._r.append(fldChar)

        self.doc.add_page_break()

    def add_list_of_figures(self):
        """Add List of Figures placeholder."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("LIST OF FIGURES")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = DEFAULT_FONT
        run.font.underline = True

        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        run = p.add_run("[This page will list all figures with their page numbers. Update after inserting actual diagrams.]")
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.size = Pt(11)
        run.font.name = DEFAULT_FONT

        self.doc.add_page_break()

    def add_list_of_tables(self):
        """Add List of Tables placeholder."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("LIST OF TABLES")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = DEFAULT_FONT
        run.font.underline = True

        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        run = p.add_run("[This page will list all tables with their page numbers.]")
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.size = Pt(11)
        run.font.name = DEFAULT_FONT

        self.doc.add_page_break()

    def add_chapter(self, chapter_number: int, title: str, content: str):
        """
        Add a complete chapter with parsed content.
        Handles headings, code blocks, diagrams, tables, and body text.
        """
        # Chapter heading
        heading = self.doc.add_heading(
            f"Chapter {chapter_number}: {title}", level=1
        )

        self.doc.add_paragraph()  # spacing after chapter title

        # Parse and add content
        self._parse_and_add_content(content, chapter_number)

        self.doc.add_page_break()

    def _parse_and_add_content(self, content: str, chapter_number: int):
        """Parse markdown-like content and add to document with formatting."""
        lines = content.split("\n")
        i = 0
        in_code_block = False
        code_buffer = []
        code_lang = ""

        while i < len(lines):
            line = lines[i]

            # Code block start/end
            if line.strip().startswith("```"):
                if in_code_block:
                    # End code block
                    self._add_code_block("\n".join(code_buffer), code_lang)
                    code_buffer = []
                    in_code_block = False
                else:
                    # Start code block
                    in_code_block = True
                    code_lang = line.strip().replace("```", "").strip()
                i += 1
                continue

            if in_code_block:
                code_buffer.append(line)
                i += 1
                continue

            # Diagram placeholder
            if line.strip().startswith("[DIAGRAM:"):
                diagram_text = line.strip()
                # Handle multi-line diagram descriptions
                while not diagram_text.endswith("]") and i + 1 < len(lines):
                    i += 1
                    diagram_text += " " + lines[i].strip()
                self._add_diagram_placeholder(diagram_text, chapter_number)
                i += 1
                continue

            # Table detection (markdown tables)
            if "|" in line and i + 1 < len(lines) and "|" in lines[i + 1]:
                table_lines = []
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                self._add_table_from_markdown(table_lines, chapter_number)
                continue

            # Headings
            if line.strip().startswith("### "):
                heading_text = line.strip().replace("### ", "", 1)
                self.doc.add_heading(heading_text, level=3)
                i += 1
                continue

            if line.strip().startswith("## "):
                heading_text = line.strip().replace("## ", "", 1)
                self.doc.add_heading(heading_text, level=2)
                i += 1
                continue

            # Bullet points
            if line.strip().startswith("- ") or line.strip().startswith("* "):
                bullet_text = line.strip()[2:]
                p = self.doc.add_paragraph(bullet_text, style="List Bullet")
                for run in p.runs:
                    run.font.size = Pt(BODY_SIZE)
                    run.font.name = DEFAULT_FONT
                i += 1
                continue

            # Numbered lists
            numbered = re.match(r"^\s*(\d+)\.\s+(.+)", line)
            if numbered:
                p = self.doc.add_paragraph(numbered.group(2), style="List Number")
                for run in p.runs:
                    run.font.size = Pt(BODY_SIZE)
                    run.font.name = DEFAULT_FONT
                i += 1
                continue

            # Regular paragraph (skip empty lines)
            if line.strip():
                self._add_formatted_text(line.strip())

            i += 1

    def _add_formatted_text(self, text: str):
        """Add a paragraph with basic formatting (bold, italic)."""
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = LINE_SPACING

        # Handle bold and italic inline
        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                run = p.add_run(part)

            run.font.size = Pt(BODY_SIZE)
            run.font.name = DEFAULT_FONT

    def _add_code_block(self, code: str, language: str = ""):
        """Add a formatted code block with gray background."""
        # Add language label if present
        if language:
            p = self.doc.add_paragraph()
            run = p.add_run(f"Code ({language}):")
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.name = DEFAULT_FONT
            run.font.color.rgb = RGBColor(80, 80, 80)

        # Create a bordered box using a single-cell table
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell = table.rows[0].cells[0]
        # Set gray background
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

        # Add code text
        cell.paragraphs[0].clear()
        for i, code_line in enumerate(code.split("\n")):
            if i == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            run = p.add_run(code_line)
            run.font.name = CODE_FONT
            run.font.size = Pt(CODE_SIZE)
            run.font.color.rgb = RGBColor(30, 30, 30)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

        # Set table borders
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f"</w:tblBorders>"
        )
        tblPr.append(borders)

        self.doc.add_paragraph()  # spacing after code block

    def _add_diagram_placeholder(self, text: str, chapter_number: int):
        """Add a bordered placeholder box for diagrams."""
        self.figure_count += 1

        # Extract diagram description
        desc = text.replace("[DIAGRAM:", "").replace("]", "").strip()
        parts = desc.split(" - ", 1)
        diagram_type = parts[0].strip() if parts else "Diagram"
        description = parts[1].strip() if len(parts) > 1 else desc

        # Create bordered box
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell = table.rows[0].cells[0]

        # Set border
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="333333"/>'
            f'  <w:left w:val="single" w:sz="8" w:space="0" w:color="333333"/>'
            f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="333333"/>'
            f'  <w:right w:val="single" w:sz="8" w:space="0" w:color="333333"/>'
            f"</w:tblBorders>"
        )
        tblPr.append(borders)

        # Set cell width to simulate a diagram area
        cell_width = Inches(5.5)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(cell_width.emu / 914400 * 1440)}" w:type="dxa"/>')
        tcPr.append(tcW)

        # Add empty lines for space
        cell.paragraphs[0].clear()
        for _ in range(8):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

        # Center text
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[Insert {diagram_type} Here]")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.italic = True
        run.font.name = DEFAULT_FONT

        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(description)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.font.name = DEFAULT_FONT

        for _ in range(8):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

        # Figure caption
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Figure {chapter_number}.{self.figure_count}: {diagram_type}")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.name = DEFAULT_FONT

        self.doc.add_paragraph()

    def _add_table_from_markdown(self, table_lines: list, chapter_number: int):
        """Parse markdown table and add a formatted Word table."""
        self.table_count += 1

        # Parse headers and rows
        rows_data = []
        for line in table_lines:
            # Skip separator lines (|---|---|)
            if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if cells:
                rows_data.append(cells)

        if not rows_data:
            return

        num_cols = max(len(row) for row in rows_data)
        table = self.doc.add_table(rows=len(rows_data), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Style the table
        for i, row_data in enumerate(rows_data):
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = table.rows[i].cells[j]
                    cell.paragraphs[0].clear()
                    p = cell.paragraphs[0]
                    run = p.add_run(cell_text)
                    run.font.size = Pt(10)
                    run.font.name = DEFAULT_FONT

                    if i == 0:  # Header row
                        run.bold = True
                        # Gray background for header
                        shading_elm = parse_xml(
                            f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>'
                        )
                        cell._tc.get_or_add_tcPr().append(shading_elm)

        # Add borders
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
            f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
            f"</w:tblBorders>"
        )
        tblPr.append(borders)

        # Table caption
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Table {chapter_number}.{self.table_count}: Data Table")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.name = DEFAULT_FONT

        self.doc.add_paragraph()

    def add_references(self, content: str = ""):
        """Add references/bibliography section."""
        self.doc.add_heading("References", level=1)
        self.doc.add_paragraph()

        if content:
            self._add_formatted_text(content)
        else:
            p = self.doc.add_paragraph()
            run = p.add_run("[References will be listed here in IEEE/APA format]")
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.name = DEFAULT_FONT

        self.doc.add_page_break()

    def add_appendix(self, content: str = ""):
        """Add appendix section."""
        self.doc.add_heading("Appendix", level=1)
        self.doc.add_paragraph()

        if content:
            self._parse_and_add_content(content, 0)
        else:
            p = self.doc.add_paragraph()
            run = p.add_run("[Appendix - Additional source code, screenshots, etc.]")
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.name = DEFAULT_FONT

    def add_page_numbers(self):
        """Add page numbers to the footer."""
        for section in self.doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add PAGE field
            run = p.add_run()
            fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run._r.append(fldChar1)

            run = p.add_run()
            instrText = parse_xml(
                f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
            )
            run._r.append(instrText)

            run = p.add_run()
            fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run._r.append(fldChar2)

    def save(self) -> str:
        """Save the document and return the file path."""
        self.add_page_numbers()
        filepath = os.path.join(OUTPUT_DIR, self.filename)
        self.doc.save(filepath)
        return filepath
